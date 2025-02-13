from docx import Document
from docx.shared import Pt


def align_chords_to_phrase(phrase: dict, chord_segments: list) -> str:
    """
    Given a phrase (with 'start', 'end', and 'text') and a list of chord segments (with 'start' and 'chord')
    that fall within the phrase interval, returns a string (of the same length as the phrase text) with chord labels
    positioned proportionally, ensuring at least one space between chords.

    :param phrase: Dictionary with keys 'start', 'end', and 'text'.
    :param chord_segments: List of dictionaries with keys 'start' and 'chord'.
    :return: A string representing the chord line.
    """
    text = phrase["text"]
    duration = phrase["end"] - phrase["start"]
    if duration <= 0:
        duration = 1
    L = len(text)
    chord_line = [" "] * L
    last_pos = -2
    for seg in chord_segments:
        seg_start = seg["start"]
        if seg_start < phrase["start"] or seg_start > phrase["end"]:
            continue
        rel = (seg_start - phrase["start"]) / duration
        pos = int(rel * (L - 1))
        if pos <= last_pos + 1:
            pos = last_pos + 2
        chord_label = seg["chord"]
        for j, char in enumerate(chord_label):
            index = pos + j
            if index < L:
                chord_line[index] = char
        last_pos = pos + len(chord_label)

    return "".join(chord_line)


def generate_word_document_chord_lyrics(
    phrase_chord_data: list, output_path: str
) -> None:
    """
    Generates a Word document where each phrase is displayed in two lines:
      - The first line shows the chords aligned according to their times.
      - The second line shows the transcribed text (lyrics).
    A blank line is inserted between phrases.

    :param phrase_chord_data: List of dictionaries, each with 'lyric' and 'chord_line'.
    :param output_path: Path to save the document.
    """
    document = Document()
    document.add_heading("Processed Music", level=1)
    for entry in phrase_chord_data:
        p = document.add_paragraph()
        run = p.add_run(entry["chord_line"] + "\n" + entry["lyric"])
        run.font.name = "Courier New"
        run.font.size = Pt(16)
        p.add_run("\n")
    document.save(output_path)
